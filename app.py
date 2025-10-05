"""
Advanced Resume Screening System (Single-file)
Enhanced for production-like features:
- Section parsing (Contact, Education, Experience, Skills)
- Section-aware TF-IDF + BERT (if available) similarity
- Per-keyword explainability (term importance)
- Employment gap detection & experience timeline
- Duplicate/resume-similarity detection
- Candidate comparison & CSV/JSON export
- Model caching and graceful fallbacks
- Tunable weights via sidebar
"""

import streamlit as st
import pandas as pd
import numpy as np
import re
import io
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import warnings
warnings.filterwarnings('ignore')

# Visualization
import plotly.graph_objects as go
import plotly.express as px

# NLP & ML
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from collections import Counter
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import WordNetLemmatizer
from functools import lru_cache

import streamlit as st

st.set_page_config(
    page_title="Advanced Resume Screener +",
    page_icon="🧠",
    layout="wide"
)

# Optional heavy libs (graceful fallback)
try:
    from sentence_transformers import SentenceTransformer
    BERT_AVAILABLE = True
except Exception:
    BERT_AVAILABLE = False

try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except Exception:
        SPACY_AVAILABLE = False
except Exception:
    SPACY_AVAILABLE = False

# Document libs
try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdf_extract
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

for res in ['punkt', 'stopwords', 'wordnet', 'omw-1.4']:
    try:
        if res == 'punkt':
            nltk.data.find('tokenizers/punkt')
        else:
            nltk.data.find(f'corpora/{res}')
    except LookupError:
        nltk.download(res, quiet=True)


# -----------------------------------------------------------------------------
#  Utilities and preprocessors
# -----------------------------------------------------------------------------

def safe_lower(x: str) -> str:
    return x.lower() if isinstance(x, str) else ""

class TextPreprocessor:
    def __init__(self):
        self.lem = WordNetLemmatizer()
        try:
            self.stop_words = set(stopwords.words('english'))
            # Keep some negations and modal auxiliaries if you want preserving
            self.stop_words -= {'not', 'no', 'very', 'can', 'will', 'should'}
        except Exception:
            self.stop_words = set()
    
    def clean_text(self, text: str) -> str:
        if not text:
            return ""
        text = re.sub(r'http\S+|www\S+', ' ', text)
        text = re.sub(r'\S+@\S+', ' ', text)
        # normalize bullets and separators
        text = text.replace('\r', '\n')
        text = re.sub(r'[^A-Za-z0-9\.\,\-\n ]', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def tokenize(self, text: str) -> List[str]:
        if not text:
            return []
        tokens = word_tokenize(text.lower())
        tokens = [t for t in tokens if t.isalnum() and len(t) > 2]
        tokens = [self.lem.lemmatize(t) for t in tokens]
        tokens = [t for t in tokens if t not in self.stop_words]
        return tokens
    
    def preprocess(self, text: str) -> str:
        tokens = self.tokenize(self.clean_text(text))
        return ' '.join(tokens)
    
    def sentences(self, text: str) -> List[str]:
        if not text:
            return []
        return sent_tokenize(text)

preprocessor = TextPreprocessor()

# -----------------------------------------------------------------------------
#  Document extraction helpers
# -----------------------------------------------------------------------------

class DocumentExtractor:
    @staticmethod
    def extract_from_pdf(filelike) -> str:
        if not PDF_AVAILABLE:
            return ""
        try:
            filelike.seek(0)
            text = pdf_extract(filelike)
            if text and text.strip():
                return text
        except Exception:
            pass
        # Fallback: PyPDF2
        try:
            filelike.seek(0)
            reader = PyPDF2.PdfReader(filelike)
            pages = []
            for p in reader.pages:
                try:
                    pages.append(p.extract_text() or "")
                except Exception:
                    pass
            return "\n".join(pages)
        except Exception:
            return ""
    
    @staticmethod
    def extract_from_docx(filelike) -> str:
        if not DOCX_AVAILABLE:
            return ""
        try:
            filelike.seek(0)
            doc = Document(filelike)
            parts = [p.text for p in doc.paragraphs if p.text]
            # tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return ""
    
    @staticmethod
    def extract_from_txt(filelike) -> str:
        try:
            filelike.seek(0)
            raw = filelike.read()
            if isinstance(raw, bytes):
                try:
                    return raw.decode('utf-8', errors='ignore')
                except:
                    return raw.decode('latin-1', errors='ignore')
            return raw
        except Exception:
            return ""
    
    @staticmethod
    def extract_text(filelike, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return DocumentExtractor.extract_from_pdf(filelike)
        elif ext == '.docx':
            return DocumentExtractor.extract_from_docx(filelike)
        elif ext == '.txt':
            return DocumentExtractor.extract_from_txt(filelike)
        else:
            # Try reading as text
            return DocumentExtractor.extract_from_txt(filelike)

# -----------------------------------------------------------------------------
#  Resume section parsers
# -----------------------------------------------------------------------------

def find_contact_info(text: str) -> Dict:
    """Attempt to parse name/email/phone/linkedin/github"""
    res = {'emails': [], 'phones': [], 'linkedin': [], 'github': [], 'name': None}
    if not text:
        return res
    # emails
    res['emails'] = list(set(re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', text)))
    # phones (simple)
    phones = re.findall(r'(?:\+?\d{1,3}[\s-])?(?:\d{10}|\d{5}[\s-]\d{5}|\d{3}[\s-]\d{3}[\s-]\d{4})', text)
    res['phones'] = list(set([re.sub(r'\D', '', p) for p in phones if len(re.sub(r'\D', '', p)) >= 8]))
    # linkedin/github
    lk = re.findall(r'(https?://(?:www\.)?linkedin\.com/[^\s,;\n]+)', text, flags=re.I)
    gh = re.findall(r'(https?://(?:www\.)?github\.com/[^\s,;\n]+)', text, flags=re.I)
    res['linkedin'] = list(set(lk))
    res['github'] = list(set(gh))
    # simple name guess: first non-empty line that's short and contains words
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines:
        # prefer line with alphabetic and length < 40 and not containing job title keywords
        for ln in lines[:6]:
            if len(ln) < 40 and len([w for w in ln.split() if w.isalpha()]) >= 1 and 'resume' not in ln.lower():
                res['name'] = re.sub(r'[^A-Za-z \-\.]', '', ln).strip()
                break
    return res

def extract_sections(text: str) -> Dict[str, str]:
    """Heuristic-based splitting into sections (Experience, Education, Skills, Summary)"""
    sections = {'summary': '', 'experience': '', 'education': '', 'skills': '', 'other': ''}
    if not text:
        return sections
    t = text
    # Normalize headings: look for common section headers
    headings = {
        'experience': r'(experience|work experience|professional experience|employment history|work history)',
        'education': r'(education|academic|academic background|qualifications)',
        'skills': r'(skill|skills|technical skills|technical competency|technical competencies)',
        'summary': r'(summary|professional summary|profile|about|objective|career objective)'
    }
    # split by lines
    lines = t.splitlines()
    # Build an index of heading line numbers
    header_indices = {}
    for i, ln in enumerate(lines):
        for sec, pat in headings.items():
            if re.search(r'\b' + pat + r'\b', ln, flags=re.I):
                header_indices[sec] = i
    # Determine ranges by sorting indices
    if header_indices:
        sorted_secs = sorted(header_indices.items(), key=lambda x: x[1])
        # for each section, take text until next known section
        for idx, (sec, start) in enumerate(sorted_secs):
            end = len(lines)
            if idx + 1 < len(sorted_secs):
                end = sorted_secs[idx + 1][1]
            sections[sec] = "\n".join(lines[start+1:end]).strip()
    # Fallback: rough heuristics if certain sections empty
    if not sections['skills']:
        # look for skills inline: comma separated lists with known tech
        skills_guess = re.findall(r'([A-Za-z\+\#\.\-]{2,}\s*(?:,|\||·|\band\b)\s*[A-Za-z0-9\+\#\.\-]{2,})', t)
        if skills_guess:
            sections['skills'] = ', '.join(skills_guess[:3])
    if not any(sections.values()):
        # fallback whole text as 'other'
        sections['other'] = t
    return sections

def parse_experience_entries(experience_text: str) -> List[Dict]:
    """Attempt to parse experience bullets into (role, company, start, end, desc)"""
    entries = []
    if not experience_text:
        return entries
    # Split by common bullet markers or double newlines
    parts = re.split(r'\n\s*(?:-|\u2022|\*)\s+', experience_text)
    if len(parts) == 1:
        parts = re.split(r'\n{2,}', experience_text)
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # find a year range like 2019 - 2021, Jan 2020 — Present, 2022 to 2024
        yrs = re.findall(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\.?\s*\d{4})', p, flags=re.I)
        range_match = re.search(r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\.?\s*\d{4})\s*(?:-|–|—|to)\s*(Present|present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\.?\s*\d{4})', p)
        start, end = None, None
        if range_match:
            start = range_match.group(1)
            end = range_match.group(2)
        # role/company attempt: lines starting with role at company
        lines = p.splitlines()
        role_company = lines[0] if lines else p
        # split by ' at ' or ',' heuristics
        rc_split = re.split(r'\s+at\s+|,', role_company, maxsplit=1)
        role = rc_split[0].strip()
        company = rc_split[1].strip() if len(rc_split) > 1 else ''
        desc = " ".join(lines[1:]).strip() if len(lines) > 1 else ''
        entries.append({'role': role, 'company': company, 'start': start, 'end': end, 'description': desc, 'raw': p})
    return entries

# -----------------------------------------------------------------------------
#  Keyword and skills extraction
# -----------------------------------------------------------------------------

DEFAULT_TECH_KEYWORDS = {
    'python','java','javascript','typescript','c++','c#','ruby','php','swift','kotlin',
    'react','angular','vue','node','django','flask','spring','express',
    'sql','mysql','postgresql','mongodb','redis','elasticsearch',
    'aws','azure','gcp','docker','kubernetes','jenkins','gitlab','github',
    'machine learning','deep learning','nlp','computer vision','tensorflow','pytorch',
    'data science','analytics','tableau','power bi','excel','agile','scrum','jira','git','ci/cd','devops',
    'api','rest','graphql','microservices','cloud','html','css','sass','bootstrap','tailwind'
}

class KeywordExtractor:
    def __init__(self, tech_vocab: set = None):
        self.tech_vocab = set(x.lower() for x in (tech_vocab or DEFAULT_TECH_KEYWORDS))
        self.pre = preprocessor
    
    def extract_by_frequency(self, text: str, top_n: int = 30) -> List[Tuple[str, int]]:
        tokens = self.pre.tokenize(self.pre.clean_text(text))
        freq = Counter(tokens)
        return freq.most_common(top_n)
    
    def extract_technical_skills(self, text: str) -> List[str]:
        text_lower = text.lower()
        found = []
        for skill in self.tech_vocab:
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found.append(skill)
        return list(sorted(set(found)))
    
    def noun_phrases(self, text: str, top_n: int = 20) -> List[str]:
        if not SPACY_AVAILABLE:
            return []
        try:
            doc = nlp(text[:200000])
            phrases = [chunk.text.lower().strip() for chunk in doc.noun_chunks if 1 <= len(chunk.text.split()) <= 4]
            freq = Counter(phrases)
            return [p for p, _ in freq.most_common(top_n)]
        except Exception:
            return []

keyword_extractor = KeywordExtractor()

# -----------------------------------------------------------------------------
#  Similarity calculators
# -----------------------------------------------------------------------------

class SimilarityCalculator:
    def __init__(self, use_bert: bool = BERT_AVAILABLE):
        self.pre = preprocessor
        self.tfidf_vectorizer = TfidfVectorizer(max_features=1000, ngram_range=(1,2))
        self.bert_model = None
        if use_bert and BERT_AVAILABLE:
            try:
                # lazy load
                self.bert_model = SentenceTransformer('all-MiniLM-L6-v2')
            except Exception:
                self.bert_model = None
    
    def tfidf_similarity(self, text1: str, text2: str) -> float:
        a = self.pre.preprocess(text1)
        b = self.pre.preprocess(text2)
        if not a or not b:
            return 0.0
        vecs = self.tfidf_vectorizer.fit_transform([a,b])
        try:
            sim = cosine_similarity(vecs[0:1], vecs[1:2])[0][0]
            return float(np.clip(sim, 0.0, 1.0))
        except Exception:
            return 0.0
    
    def bert_similarity(self, text1: str, text2: str) -> float:
        if not self.bert_model:
            return 0.0
        try:
            # encoding longer docs by chunks can be added; keep short for speed
            emb1 = self.bert_model.encode(text1[:2000], convert_to_tensor=False)
            emb2 = self.bert_model.encode(text2[:2000], convert_to_tensor=False)
            sim = np.dot(emb1, emb2) / (np.linalg.norm(emb1) * np.linalg.norm(emb2) + 1e-10)
            return float(np.clip(sim, -1.0, 1.0))
        except Exception:
            return 0.0
    
    def keyword_overlap(self, jd_keywords: List[str], resume_text: str) -> Dict:
        resume_lower = resume_text.lower()
        found = [kw for kw in jd_keywords if kw.lower() in resume_lower]
        missing = [kw for kw in jd_keywords if kw.lower() not in resume_lower]
        coverage = (len(found) / len(jd_keywords) * 100) if jd_keywords else 0.0
        return {'found': found, 'missing': missing, 'coverage': coverage}

sim_calc = SimilarityCalculator()

# -----------------------------------------------------------------------------
#  Resume analyzer (main scoring logic)
# -----------------------------------------------------------------------------

def extract_experience_years(text: str) -> float:
    """Simple heuristic to estimate years of experience from text"""
    # Look for patterns like '4 years', '5+ years'
    matches = re.findall(r'(\d{1,2})\+?\s*(?:years?|yrs?)', text.lower())
    nums = [int(m) for m in matches if m.isdigit()]
    if nums:
        return max(nums)
    # fallback: infer from date ranges in experience section
    years = re.findall(r'(19|20)\d{2}', text)
    if len(years) >= 2:
        years_int = [int(y) for y in years]
        return max(years_int) - min(years_int)
    return 0.0

def detect_employment_gaps(entries: List[Dict]) -> List[Tuple[str, str]]:
    """Return list of (gap_start, gap_end) if gaps detected - heuristic"""
    # convert start/end to years if possible
    def parse_year(x):
        if not x:
            return None
        yy = re.search(r'(19|20)\d{2}', x)
        if yy:
            return int(yy.group(0))
        return None
    years = []
    for e in entries:
        s = parse_year(e.get('start'))
        e_ = parse_year(e.get('end')) if e.get('end') and e.get('end').lower() not in ('present','current') else datetime.now().year
        if s:
            years.append((s, e_ or s))
    if not years:
        return []
    # sort and find gaps
    years = sorted(years, key=lambda x: x[0])
    gaps = []
    for i in range(len(years)-1):
        if years[i+1][0] - years[i][1] >= 2:  # gap of 2+ years
            gaps.append((str(years[i][1]), str(years[i+1][0])))
    return gaps

def explain_keyword_contribution(jd_text: str, resume_text: str, top_n: int = 15) -> List[Tuple[str, float]]:
    """Return top terms from JD and their TF-IDF weights in resume (approximated)"""
    # Use TF-IDF computed on JD+resume and get tfidf for JD terms when mapped to resume
    jd_proc = preprocessor.preprocess(jd_text)
    res_proc = preprocessor.preprocess(resume_text)
    if not jd_proc or not res_proc:
        return []
    vect = TfidfVectorizer(max_features=1000, ngram_range=(1,2))
    tfidf = vect.fit_transform([jd_proc, res_proc])
    feature_names = vect.get_feature_names_out()
    jd_vec = tfidf[0].toarray().ravel()
    res_vec = tfidf[1].toarray().ravel()
    # pick top JD terms by jd_vec weight, but show their presence weight in resume (res_vec)
    jd_indices = np.argsort(jd_vec)[::-1][:top_n]
    results = []
    for idx in jd_indices:
        term = feature_names[idx]
        jd_w = float(jd_vec[idx])
        res_w = float(res_vec[idx])
        results.append((term, res_w))
    return results

def score_resume(jd_text: str, resume_text: str, weights: Dict[str, float]) -> Dict:
    # extract sections & keywords
    sections = extract_sections(resume_text)
    kd = keyword_extractor.extract_by_frequency(jd_text, top_n=40)
    jd_keywords = [k for k, _ in kd]
    # Compute similarities across sections (global, summary, experience, skills)
    global_tfidf = sim_calc.tfidf_similarity(jd_text, resume_text)
    summary_tfidf = sim_calc.tfidf_similarity(jd_text, sections.get('summary',''))
    exp_tfidf = sim_calc.tfidf_similarity(jd_text, sections.get('experience',''))
    skills_tfidf = sim_calc.tfidf_similarity(jd_text, sections.get('skills',''))
    bert_sim = sim_calc.bert_similarity(jd_text, resume_text) if sim_calc.bert_model else 0.0
    # keyword overlap
    kw_info = sim_calc.keyword_overlap(jd_keywords, resume_text)
    # tech skills
    jd_tech = keyword_extractor.extract_technical_skills(jd_text)
    res_tech = keyword_extractor.extract_technical_skills(resume_text)
    tech_match_count = len(set(jd_tech) & set(res_tech))
    tech_score = (tech_match_count / len(jd_tech) * 100) if jd_tech else 0
    # experience
    exp_years = extract_experience_years(resume_text)
    exp_score = min(exp_years / 10 * 100, 100)
    # explain terms
    term_contrib = explain_keyword_contribution(jd_text, resume_text, top_n=20)
    # experience entries & gaps
    exp_entries = parse_experience_entries(sections.get('experience',''))
    gaps = detect_employment_gaps(exp_entries)
    # duplicate / fingerprint: quick hash of skills + top terms
    fingerprint = "_".join(sorted(res_tech)) + "__" + "_".join([t for t,_ in kd[:5]])
    # Normalized combining (convert similarities to 0-100)
    global_tfidf_pct = global_tfidf * 100
    bert_pct = (bert_sim + 1)/2 * 100 if bert_sim != 0 else 0  # convert [-1,1] to [0,100]
    # Weighted final score
    # weights dict should sum to 1 (we normalize to be safe)
    wsum = sum(weights.values()) if weights else 1.0
    if wsum == 0:
        wsum = 1.0
    normalized_weights = {k: v/wsum for k, v in weights.items()}
    final = (
        normalized_weights.get('bert',0) * bert_pct +
        normalized_weights.get('tfidf',0) * global_tfidf_pct +
        normalized_weights.get('section_exp',0) * (exp_tfidf*100) +
        normalized_weights.get('section_skills',0) * (skills_tfidf*100) +
        normalized_weights.get('keywords',0) * kw_info['coverage'] +
        normalized_weights.get('tech',0) * tech_score +
        normalized_weights.get('experience',0) * exp_score
    )
    final = float(np.clip(final, 0.0, 100.0))
    # Recommendation thresholds
    if final >= 75:
        rec = 'Strong Match'
        color = 'green'
    elif final >= 55:
        rec = 'Good Match'
        color = 'blue'
    elif final >= 35:
        rec = 'Consider'
        color = 'orange'
    else:
        rec = 'Weak'
        color = 'red'
    return {
        'final_score': round(final,2),
        'tfidf_score': round(global_tfidf_pct,2),
        'bert_score': round(bert_pct,2),
        'summary_tfidf': round(summary_tfidf*100,2),
        'experience_tfidf': round(exp_tfidf*100,2),
        'skills_tfidf': round(skills_tfidf*100,2),
        'keyword_coverage': round(kw_info['coverage'],2),
        'keywords_found': kw_info['found'],
        'keywords_missing': kw_info['missing'],
        'jd_tech_skills': jd_tech,
        'resume_tech_skills': res_tech,
        'tech_score': round(tech_score,2),
        'experience_years': exp_years,
        'experience_score': round(exp_score,2),
        'term_contrib': term_contrib,
        'experience_entries': exp_entries,
        'employment_gaps': gaps,
        'fingerprint': fingerprint,
        'recommendation': rec,
        'recommendation_color': color
    }

# -----------------------------------------------------------------------------
#  Streamlit UI: helpers for charts & gauge
# -----------------------------------------------------------------------------

def create_score_gauge(score: float, title: str) -> go.Figure:
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': title},
        gauge={'axis': {'range': [0,100]},
               'bar': {'color': "darkblue"},
               'steps': [
                   {'range': [0,35], 'color': "lightgray"},
                   {'range': [35,55], 'color': "orange"},
                   {'range': [55,75], 'color': "lightblue"},
                   {'range': [75,100], 'color': "lightgreen"}
               ],
               'threshold': {
                   'line': {'color': "red", 'width': 4},
                   'thickness': 0.75,
                   'value': 75
               }
        }
    ))
    fig.update_layout(height=280, margin=dict(l=10,r=10,t=40,b=10))
    return fig

# -----------------------------------------------------------------------------
#  App main
# -----------------------------------------------------------------------------

# st.set_page_config(page_title="Advanced Resume Screener +", page_icon="🧠", layout="wide")
st.title("🧠 Advanced Resume Screening System — Enhanced")

# sidebar config
with st.sidebar:
    st.header("Configuration & Weights")
    use_bert = st.checkbox("Use BERT semantic similarity (heavy)", value=BERT_AVAILABLE and sim_calc.bert_model is not None)
    # Weight sliders
    w_tfidf = st.slider("TF-IDF weight", 0.0, 1.0, 0.25)
    w_bert = st.slider("BERT weight", 0.0, 1.0, 0.35 if BERT_AVAILABLE else 0.0)
    w_section_exp = st.slider("Experience section weight", 0.0, 1.0, 0.10)
    w_section_skills = st.slider("Skills section weight", 0.0, 1.0, 0.10)
    w_keywords = st.slider("Keyword coverage weight", 0.0, 1.0, 0.10)
    w_tech = st.slider("Tech skill match weight", 0.0, 1.0, 0.05)
    w_experience = st.slider("Experience years weight", 0.0, 1.0, 0.05)
    threshold = st.slider("Show candidates with score ≥", 0, 100, 35)
    st.markdown("---")
    st.write("Model status:")
    st.write(f"Sentence-BERT available: {'✅' if BERT_AVAILABLE else '❌'}")
    st.write(f"SpaCy available: {'✅' if SPACY_AVAILABLE else '❌'}")
    st.write(f"PDF support: {'✅' if PDF_AVAILABLE else '❌'}")
    st.markdown("---")
    st.caption("Tip: Set weights to reflect your hiring priorities. Total weight will be normalized.")

weights = {
    'tfidf': w_tfidf,
    'bert': w_bert if use_bert else 0.0,
    'section_exp': w_section_exp,
    'section_skills': w_section_skills,
    'keywords': w_keywords,
    'tech': w_tech,
    'experience': w_experience
}

# main UI layout: tabs
tab1, tab2, tab3 = st.tabs(["📤 Upload & Configure", "📊 Results Dashboard", "🔎 Compare & Export"])

# Session-state initialization
if 'results' not in st.session_state:
    st.session_state.results = []
if 'jd_text' not in st.session_state:
    st.session_state.jd_text = ""
if 'last_run' not in st.session_state:
    st.session_state.last_run = None

# Upload & JD input
with tab1:
    st.subheader("1) Job Description (JD)")
    jd_input_method = st.radio("JD input method:", ("Type/Paste Text", "Upload File (pdf/docx/txt)"), horizontal=True)
    jd_text = ""
    if jd_input_method == "Type/Paste Text":
        jd_text = st.text_area("Paste full job description (include responsibilities, must-have skills):",
                               height=250, value=st.session_state.jd_text or "")
    else:
        jd_file = st.file_uploader("Upload JD file", type=['pdf','docx','txt'], key="jd_file")
        if jd_file:
            with st.spinner("Extracting JD text..."):
                jd_text = DocumentExtractor.extract_text(jd_file, jd_file.name)
                st.success(f"Extracted {len(jd_text.split())} words from {jd_file.name}")
    if jd_text:
        st.session_state.jd_text = jd_text
        jd_keywords = keyword_extractor.extract_by_frequency(jd_text, top_n=30)
        col1, col2, col3 = st.columns(3)
        col1.metric("Words in JD", len(jd_text.split()))
        col2.metric("Top terms", len(jd_keywords))
        col3.metric("Detected tech skills", len(keyword_extractor.extract_technical_skills(jd_text)))
        with st.expander("Preview JD and extracted keywords"):
            st.text_area("Job Description", jd_text, height=180)
            df = pd.DataFrame(jd_keywords, columns=['term','freq'])
            st.dataframe(df.head(30))
    st.markdown("---")
    st.subheader("2) Upload Resumes")
    uploaded = st.file_uploader("Upload multiple resumes", accept_multiple_files=True, type=['pdf','docx','txt'])
    if uploaded:
        st.success(f"{len(uploaded)} files selected")
        with st.expander("Show uploaded files"):
            for f in uploaded:
                st.write(f"{f.name} — {f.size/1024:.1f} KB")
    st.markdown("---")
    st.write("3) Action")
    col1, col2 = st.columns([1,1])
    with col1:
        run_btn = st.button("🚀 Analyze Resumes", disabled=(not uploaded or not jd_text))
    with col2:
        clear_btn = st.button("🧹 Clear previous results")
    if clear_btn:
        st.session_state.results = []
        st.success("Cleared previous results.")

    if run_btn:
        with st.spinner("Processing resumes..."):
            results = []
            duplicates = {}
            for f in uploaded:
                text = DocumentExtractor.extract_text(f, f.name)
                if not text or not text.strip():
                    st.warning(f"Could not extract text from {f.name}")
                    continue
                # basic contact & sections
                contact = find_contact_info(text)
                sections = extract_sections(text)
                # run scoring
                sim_calc_local = SimilarityCalculator(use_bert=use_bert)
                # attach BERT model to sim_calc if use_bert and model available
                # (here SimilarityCalculator handles it)
                analysis = score_resume(st.session_state.jd_text, text, weights)
                # add metadata
                analysis.update({
                    'name': contact.get('name') or f.name,
                    'file_name': f.name,
                    'emails': contact.get('emails',[]),
                    'phones': contact.get('phones',[]),
                    'linkedin': contact.get('linkedin',[]),
                    'github': contact.get('github',[]),
                    'sections': sections,
                    'raw_text': text
                })
                # duplicate detection via fingerprint
                fingerprint = analysis.get('fingerprint','')
                duplicates.setdefault(fingerprint, []).append(analysis['file_name'])
                results.append(analysis)
            # mark duplicates in results metadata
            for r in results:
                r['duplicates'] = duplicates.get(r.get('fingerprint',''), [])
            # sort
            results = sorted(results, key=lambda x: x['final_score'], reverse=True)
            st.session_state.results = results
            st.session_state.last_run = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.success(f"Analyzed {len(results)} resumes — completed at {st.session_state.last_run}")
            st.balloons()

# Results Dashboard
with tab2:
    st.subheader("Results Dashboard")
    if not st.session_state.results:
        st.info("Upload JD and resumes in the 'Upload & Configure' tab and run analysis.")
    else:
        results = st.session_state.results
        df = pd.DataFrame([{
            'Rank': idx+1,
            'Candidate': r.get('name') or r.get('file_name'),
            'File': r.get('file_name'),
            'Final Score': r['final_score'],
            'Recommendation': r['recommendation'],
            'TF-IDF': r['tfidf_score'],
            'BERT': r['bert_score'],
            'Keyword Coverage': r['keyword_coverage'],
            'Tech Score': r['tech_score'],
            'Experience (yrs)': r['experience_years']
        } for idx, r in enumerate(results)])
        # filter by threshold
        filtered_df = df[df['Final Score'] >= threshold].reset_index(drop=True)
        col1, col2 = st.columns([2,1])
        with col1:
            st.dataframe(filtered_df, use_container_width=True, height=400)
        with col2:
            st.metric("Total candidates", len(df))
            st.metric("Above threshold", len(filtered_df))
            st.metric("Last run", st.session_state.last_run or "N/A")
            if st.button("🔄 Re-sort by tech score"):
                df = df.sort_values('Tech Score', ascending=False)
        st.markdown("---")
        st.subheader("Detailed candidate view")
        candidate = st.selectbox("Choose candidate", options=[r.get('name') or r.get('file_name') for r in results])
        if candidate:
            cand = next(r for r in results if (r.get('name') == candidate or r.get('file_name') == candidate))
            c1, c2, c3 = st.columns([1,1,1])
            with c1:
                st.write(f"**Name:** {cand.get('name')}")
                st.write(f"**File:** {cand.get('file_name')}")
                st.write(f"**Emails:** {', '.join(cand.get('emails',[])) or 'N/A'}")
                st.write(f"**Phones:** {', '.join(cand.get('phones',[])) or 'N/A'}")
                st.write(f"**LinkedIn:** {', '.join(cand.get('linkedin',[])) or 'N/A'}")
            with c2:
                st.write("**Recommendation**")
                st.markdown(f"<div style='color:{cand['recommendation_color']}'>{cand['recommendation']} — {cand['final_score']}%</div>", unsafe_allow_html=True)
                st.plotly_chart(create_score_gauge(cand['final_score'], "Final Score"), use_container_width=True)
            with c3:
                st.write("**Key metrics**")
                st.metric("TF-IDF", f"{cand['tfidf_score']}%")
                st.metric("BERT", f"{cand['bert_score']}%")
                st.metric("Keyword coverage", f"{cand['keyword_coverage']}%")
                st.metric("Tech match", f"{cand['tech_score']}%")
            st.markdown("---")
            st.subheader("Top matched keywords & explainability")
            top_terms = cand['term_contrib'][:15]
            if top_terms:
                for term, score in top_terms:
                    st.write(f"• **{term}** — presence weight in resume: {score:.4f}")
            else:
                st.write("No explainable terms available.")
            st.markdown("---")
            st.subheader("Experience parsing")
            st.write(f"Detected experience years: **{cand['experience_years']}**")
            entries = cand.get('experience_entries',[])
            if entries:
                for e in entries:
                    st.write(f"- **{e.get('role','')}** at **{e.get('company','')}** ({e.get('start') or 'N/A'} — {e.get('end') or 'N/A'})")
                    if e.get('description'):
                        with st.expander("Job description"):
                            st.write(e.get('description'))
            else:
                st.write("No structured experience entries parsed.")
            if cand.get('employment_gaps'):
                st.warning(f"Employment gaps detected: {cand.get('employment_gaps')}")
            st.markdown("---")
            st.subheader("Detected technical skills")
            st.write(", ".join(cand.get('resume_tech_skills',[])) or "None detected")
            st.markdown("---")
            st.subheader("Matching sentences from resume (highest keyword matches)")
            # naive matching sentences
            resume_text = cand.get('raw_text','')
            sentences = preprocessor.sentences(resume_text)
            jd_top_terms = [k for k,_ in keyword_extractor.extract_by_frequency(st.session_state.jd_text, top_n=20)]
            matches = []
            for s in sentences:
                lower = s.lower()
                found = [t for t in jd_top_terms if t.lower() in lower]
                if found:
                    matches.append({'sentence': s.strip(), 'count': len(found), 'keywords': found})
            matches = sorted(matches, key=lambda x: x['count'], reverse=True)[:8]
            if matches:
                for m in matches:
                    with st.expander(f"Score: {m['count']} keywords"):
                        st.write(m['sentence'])
                        st.write("Keywords: " + ", ".join(m['keywords']))
            else:
                st.write("No highly matching sentences found.")
            st.markdown("---")
            st.download_button("📥 Download candidate JSON", data=io.BytesIO(str(cand).encode('utf-8')), file_name=f"{cand.get('file_name')}_analysis.json")

# Comparison & Export
with tab3:
    st.subheader("Compare Candidates & Export")
    results = st.session_state.results
    if not results:
        st.info("No results — run analysis first.")
    else:
        names = [r.get('name') or r.get('file_name') for r in results]
        selected = st.multiselect("Select candidates to compare (2-4)", options=names, default=names[:2])
        if selected:
            chosen = [r for r in results if (r.get('name') in selected or r.get('file_name') in selected)]
            if len(chosen) >= 1:
                # Radar chart style: convert metrics to 0-1
                metrics = ['tfidf_score','bert_score','keyword_coverage','tech_score','experience_score','final_score']
                comp_df = pd.DataFrame([{ 'candidate': c.get('name') or c.get('file_name'), **{m: c.get(m,0) for m in metrics}} for c in chosen])
                df = comp_df.set_index('candidate').T

                # Make duplicate column names unique (e.g., Sowjanya Thota_0, Sowjanya Thota_1, etc.)
                df.columns = [f"{col}_{i}" if list(df.columns).count(col) > 1 else col 
                            for i, col in enumerate(df.columns)]

                st.dataframe(df)

                fig = px.bar(comp_df.melt(id_vars='candidate', value_vars=metrics), x='candidate', y='value', color='variable', barmode='group', title="Metric comparison")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("Select at least one candidate.")
        # Export all results CSV/JSON
        if st.button("Export all results (CSV)"):
            df_export = pd.DataFrame(results)
            csv_bytes = df_export.to_csv(index=False).encode('utf-8')
            st.download_button("Download CSV", data=csv_bytes, file_name=f"resume_screening_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv", mime="text/csv")
        if st.button("Export all results (JSON)"):
            import json
            json_bytes = json.dumps(results, default=str, indent=2).encode('utf-8')
            st.download_button("Download JSON", data=io.BytesIO(json_bytes), file_name=f"resume_screening_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")

# Footer / tips
st.markdown("---")
st.caption("Notes: For production readiness you'd want persistence (DB), authentication, batching, robust PDF parsing tools (OCR for scanned PDFs), concurrency for large volumes, and unit tests. This app focuses on advanced features and explainability to help land high-paying roles by demonstrating careful engineering and ML system thinking.")
